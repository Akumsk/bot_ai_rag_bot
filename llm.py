import os
import pandas as pd
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyMuPDFLoader
from langchain.text_splitter import CharacterTextSplitter
from langchain.chains import RetrievalQA
from langchain.schema import Document  # Import the correct Document schema from LangChain
from docx import Document as Doc  # For Word documents
from io import StringIO
from tiktoken import encoding_for_model

# Set up OpenAI API key
openai_api_key = os.getenv('OPENAI_API_KEY')

# Initialize LLM
llm = ChatOpenAI(openai_api_key=openai_api_key, model_name='gpt-4o')

vector_store = None

######################################################################



def load_excel_file(file_path):
    """Load content from an Excel file as text."""
    data = pd.read_excel(file_path)
    text_data = StringIO()
    data.to_string(buf=text_data)
    return text_data.getvalue()

def load_word_file(file_path):
    """Load content from a Word file (.docx)."""
    doc = Doc(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

# Updated function to load and process PDF, Word, and Excel files and create FAISS index
# Updated function to load and process PDF, Word, and Excel files and create FAISS index
def load_and_index_documents(folder_path):
    global vector_store

    # Load and read documents (PDF, Word, Excel) from folder
    documents = []
    found_valid_file = False  # Track whether any valid documents are found

    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)

        # Handle PDF files
        if filename.endswith(".pdf"):
            loader = PyMuPDFLoader(file_path)
            docs = loader.load()
            for doc in docs:
                doc.metadata = {"source": filename}  # Attach the filename to the document metadata
                documents.append(doc)  # Add each PDF page as a document object
            found_valid_file = True  # Mark that we found a valid file

        # Handle Word files
        elif filename.endswith(".docx"):
            content = load_word_file(file_path)
            doc = Document(page_content=content, metadata={"source": filename})  # Use Langchain's Document schema
            documents.append(doc)
            found_valid_file = True  # Mark that we found a valid file

        # Handle Excel files
        elif filename.endswith(".xlsx"):
            content = load_excel_file(file_path)
            doc = Document(page_content=content, metadata={"source": filename})  # Use Langchain's Document schema
            documents.append(doc)
            found_valid_file = True  # Mark that we found a valid file

    # If no valid files were found, provide an appropriate error message
    if not found_valid_file:
        return "No valid files found in the folder. Please provide PDF, Word, or Excel files."

    # Split documents into smaller chunks
    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    split_docs = text_splitter.split_documents(documents)

    # Create embeddings using OpenAI embeddings
    embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)

    # Use FAISS from Langchain to store the document embeddings
    vector_store = FAISS.from_documents(split_docs, embeddings)
    return "Documents successfully indexed."


# Token counting function
def count_tokens_in_documents(documents):
    """Counts the total number of tokens in a list of documents."""
    tokenizer = encoding_for_model('gpt-4')  # Tokenizer for the specific model

    total_tokens = 0
    for doc in documents:
        total_tokens += len(tokenizer.encode(doc.page_content))

    return total_tokens


# Function to evaluate tokens in the context folder
def evaluate_context_token_count(folder_path, token_limit):
    """Evaluates the total token count in documents within a folder."""
    documents = []
    found_valid_file = False  # Track whether any valid documents are found

    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)

        # Handle PDF files
        if filename.endswith(".pdf"):
            loader = PyMuPDFLoader(file_path)
            docs = loader.load()
            for doc in docs:
                documents.append(doc)  # Add each PDF page as a document object
            found_valid_file = True  # Mark that we found a valid file

        # Handle Word files
        elif filename.endswith(".docx"):
            content = load_word_file(file_path)
            doc = Document(page_content=content, metadata={"source": filename})  # Use Langchain's Document schema
            documents.append(doc)
            found_valid_file = True  # Mark that we found a valid file

        # Handle Excel files
        elif filename.endswith(".xlsx"):
            content = load_excel_file(file_path)
            doc = Document(page_content=content, metadata={"source": filename})  # Use Langchain's Document schema
            documents.append(doc)
            found_valid_file = True  # Mark that we found a valid file

    # If no valid files were found, return an appropriate message
    if not found_valid_file:
        return "No valid files found in the folder."

    # Count tokens in the documents
    total_tokens = count_tokens_in_documents(documents)

    # If tokens exceed 10,000, return the warning message
    if total_tokens > token_limit:
        return "The context folder contains too many tokens. Kindly remove any unnecessary documents to proceed."
    else:
        return f"Total token count: {total_tokens}"

# Function to handle retrieving and generating response using RAG
def retrieve_and_generate(prompt: str):
    global vector_store
    if not vector_store:
        return "Please set the folder path using /path_folder and ensure documents are loaded.", None

    # Set up retriever
    retriever = vector_store.as_retriever()

    # Use Langchain's RetrievalQA Chain to get the response
    qa_chain = RetrievalQA.from_chain_type(
        llm=llm,
        retriever=retriever,
        chain_type="stuff",
        return_source_documents=True  # Ensure that source documents are returned
    )

    try:
        # Use .invoke() method instead of deprecated __call__()
        result = qa_chain.invoke({"query": prompt})

        # Ensure 'source_documents' key is present and retrieve documents
        sources = result.get("source_documents", [])
        if not sources:
            return result["result"], None  # Return response without source if no documents were retrieved

        # Extract the filenames from the source documents
        source_files = set([doc.metadata["source"] for doc in sources if "source" in doc.metadata])

        response = result["result"]
        return response, source_files

    except KeyError as e:
        # Handle unexpected keys in the result
        return f"Error: Missing key {str(e)} in response.", None

    except Exception as e:
        # Catch any other errors
        return f"An error occurred: {str(e)}", None